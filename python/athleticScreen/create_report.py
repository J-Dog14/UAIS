"""
Report generation for Athletic Screen.
Generates Word documents with charts, graphs, and analysis comparing client data to "All" reference data.

Based on cell 3 of Create Report_College.ipynb, adapted for PostgreSQL.
"""
import os
import sys
import tempfile
import re
import glob as globmod
from pathlib import Path
from datetime import date, datetime
from typing import List, Dict, Optional, Tuple

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from scipy import integrate, stats
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add python directory to path for imports
python_dir = Path(__file__).parent.parent
if str(python_dir) not in sys.path:
    sys.path.insert(0, str(python_dir))

from common.athlete_manager import get_warehouse_connection
from common.config import get_raw_paths

# Map movement types to PostgreSQL table names
MOVEMENT_TO_PG_TABLE = {
    'CMJ': 'f_athletic_screen_cmj',
    'DJ': 'f_athletic_screen_dj',
    'PPU': 'f_athletic_screen_ppu',
    'SLV': 'f_athletic_screen_slv',
    'NMT': 'f_athletic_screen_nmt'
}

# Dark theme styling for matplotlib
plt.rcParams.update({
    "figure.facecolor": "#181818",
    "axes.facecolor": "#303030",
    "axes.edgecolor": "white",
    "axes.labelcolor": "slategrey",
    "xtick.color": "lightgrey",
    "ytick.color": "lightgrey",
    "grid.color": "dimgrey",
    "text.color": "white",
})


def calculate_percentile(value: float, reference_data: np.ndarray) -> float:
    """Calculate percentile using scipy stats."""
    return stats.percentileofscore(reference_data, value)


def generate_bar_graph(variable: str, client_value: float, reference_data: np.ndarray,
                      title: str, tmpdirname: str, conn, athlete_uuid: str,
                      session_date: str, movement_type: str) -> str:
    """
    Generate bar graph comparing client value to reference distribution.
    
    Blue bars = reference distribution
    Red line = client MAX (best trial)
    Violet line = client MEAN (average of trials)
    """
    # Gather all trials for this athlete/table/date
    scores = []
    pg_table = MOVEMENT_TO_PG_TABLE.get(movement_type)
    
    if pg_table:
        try:
            # Query all trials for this athlete on this date
            with conn.cursor() as cur:
                if movement_type == 'SLV':
                    # SLV needs side column - we'll handle separately
                    cur.execute(f"""
                        SELECT {variable} FROM public.{pg_table}
                        WHERE athlete_uuid = %s AND session_date = %s
                        AND {variable} IS NOT NULL
                    """, (athlete_uuid, session_date))
                else:
                    cur.execute(f"""
                        SELECT {variable} FROM public.{pg_table}
                        WHERE athlete_uuid = %s AND session_date = %s
                        AND {variable} IS NOT NULL
                    """, (athlete_uuid, session_date))
                
                scores = [row[0] for row in cur.fetchall() if row[0] is not None]
        except Exception as e:
            print(f"Warning: Could not fetch trials for {variable}: {e}")
            scores = [client_value]
    
    if not scores:
        scores = [client_value]
    
    scores = np.asarray(scores, dtype=float)
    c_mean = scores.mean()
    
    # Choose extreme based on variable type
    if variable.upper() == "CT":  # Contact Time → use LOWEST value
        c_extreme = scores.min()
        extreme_label = "Client Min"
    else:  # Every other metric → highest value
        c_extreme = scores.max()
        extreme_label = "Client Max"
    
    perc_mean = calculate_percentile(c_mean, reference_data)
    
    # Plotting
    plt.figure(facecolor="#181818")
    ax = plt.subplot(111, facecolor="#303030")
    
    reference_plotted = False
    
    # RSI special case: 0.25-wide bars
    if variable.upper() == "RSI":
        lo = np.floor(reference_data.min() / 0.25) * 0.25
        hi = np.ceil(reference_data.max() / 0.25) * 0.25
        bins = np.arange(lo, hi + 0.25, 0.25)
        centers = bins[:-1]
        
        counts, _ = np.histogram(reference_data, bins=bins)
        ax.bar(centers, counts, width=0.25, align="edge",
               color="cornflowerblue", alpha=0.7, edgecolor="white", label="Reference")
        
        # Draw empty bins as thin outlines
        for x, h in zip(centers, counts):
            if h == 0:
                ax.bar(x, 1e-6, width=.25, align='edge',
                       color='none', edgecolor='#404040', linewidth=.5)
        
        ax.set_xticks(centers)
        ax.set_xticklabels([f"{x:.2f}" for x in centers], color="lightgrey")
        reference_plotted = True
    else:
        bins = 20
    
    # Draw reference histogram if not already plotted
    if not reference_plotted:
        ax.hist(reference_data, bins=bins, color="cornflowerblue", alpha=0.7,
                edgecolor="white", label="Reference")
    
    # Client mean/max lines
    ax.axvline(c_extreme, color="red", ls="--", lw=2, label=extreme_label)
    ax.axvline(c_mean, color="violet", ls="--", lw=2, label="Client Mean")
    
    # Cosmetics
    ax.set_xlabel(variable.replace("_", " "), color="slategrey")
    ax.set_ylabel("Frequency", color="slategrey")
    ax.tick_params(axis="x", colors="lightgrey")
    ax.tick_params(axis="y", colors="lightgrey")
    ax.grid(color="dimgrey")
    
    txt = (f"Percentile (mean): {perc_mean:.1f}%\n"
           f"Mean: {c_mean:.2f}\n"
           f"Max:  {c_extreme:.2f}")
    plt.text(0.95, 0.05, txt, ha="right", va="bottom",
             transform=ax.transAxes, color="white", fontsize=9,
             backgroundcolor="#181818")
    
    ax.legend(facecolor="black", edgecolor="grey", prop={"size": "small"}, labelcolor="grey")
    
    # Save
    out_path = os.path.join(tmpdirname, f"{variable}_histogram.png")
    plt.savefig(out_path, bbox_inches="tight", facecolor="#181818")
    plt.close()
    return out_path


def generate_scatter_plot(client_data: Dict, reference_data: pd.DataFrame,
                         x_var: str, y_var: str, title: str, tmpdirname: str) -> str:
    """Generate scatter plot comparing client to reference data."""
    plt.figure(facecolor='#181818', figsize=(6, 6))
    ax = plt.subplot(111, facecolor='#303030')
    
    # Reference data scatter
    if not reference_data.empty and x_var in reference_data.columns and y_var in reference_data.columns:
        ax.scatter(reference_data[x_var], reference_data[y_var], label='Reference',
                  alpha=0.5, color='cornflowerblue')
    
    # Client data scatter
    if isinstance(client_data[x_var], (list, np.ndarray)):
        ax.scatter(client_data[x_var], client_data[y_var], label='Client',
                  color='red', edgecolors='black', s=100)
    else:
        ax.scatter([client_data[x_var]], [client_data[y_var]], label='Client',
                  color='red', edgecolors='black', s=100)
    
    # Axis labels
    ax.set_xlabel(x_var.replace('_', ' '), color='slategrey')
    ax.set_ylabel(y_var.replace('_', ' '), color='slategrey')
    ax.tick_params(axis='x', colors='lightgrey')
    ax.tick_params(axis='y', colors='lightgrey')
    
    # Reference lines
    if not reference_data.empty:
        ax.axvline(x=np.mean(reference_data[x_var]), color='lightgrey',
                  linestyle='--', linewidth=1)
        ax.axhline(y=np.mean(reference_data[y_var]), color='lightgrey',
                  linestyle='--', linewidth=1)
    
    ax.grid(color='dimgrey')
    ax.legend(facecolor='black', edgecolor='grey', prop={'size': 'small'}, labelcolor='grey')
    
    scatter_filename = os.path.join(tmpdirname, f'{x_var}_{y_var}_scatter.png')
    plt.savefig(scatter_filename, bbox_inches='tight', facecolor='#181818')
    plt.close()
    
    return scatter_filename


def load_power_txt(txt_path: str) -> pd.Series:
    """Parse exported power file and return pandas Series."""
    power_vals = []
    in_data = False
    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
        for raw in f:
            line = raw.strip()
            if not in_data:
                if line.startswith("ITEM"):
                    in_data = True
                continue
            if not line:
                continue
            parts = re.split(r"\t+", line)
            if len(parts) < 2:
                continue
            try:
                val = float(parts[-1])
                power_vals.append(val)
            except ValueError:
                continue
    if not power_vals:
        raise ValueError(f"No power values parsed from {txt_path}")
    return pd.Series(power_vals, name="Power")


def find_power_files(movement: str, base_dir: str = None) -> List[str]:
    """Find exported power files for a movement."""
    if base_dir is None:
        try:
            raw_paths = get_raw_paths()
            base_dir = raw_paths.get('athletic_screen', r'D:/Athletic Screen 2.0/Output Files/')
        except:
            base_dir = r'D:/Athletic Screen 2.0/Output Files/'
    
    pats = [f"{movement}_Power.txt", f"{movement}*_Power.txt"]
    files = []
    for p in pats:
        files += globmod.glob(os.path.join(base_dir, p))
    return sorted(set(files))


def query_client_data(conn, athlete_uuid: str, session_date: str,
                     movement_type: str, columns: List[str]) -> Optional[Tuple]:
    """Query client data from PostgreSQL."""
    pg_table = MOVEMENT_TO_PG_TABLE.get(movement_type)
    if not pg_table:
        return None
    
    try:
        with conn.cursor() as cur:
            if movement_type == 'SLV':
                # SLV needs side handling - return both Left and Right
                cur.execute(f"""
                    SELECT {', '.join(columns)}
                    FROM public.{pg_table}
                    WHERE athlete_uuid = %s AND session_date = %s AND side = 'Left'
                    ORDER BY created_at DESC
                    LIMIT 1
                """, (athlete_uuid, session_date))
                left_data = cur.fetchone()
                
                cur.execute(f"""
                    SELECT {', '.join(columns)}
                    FROM public.{pg_table}
                    WHERE athlete_uuid = %s AND session_date = %s AND side = 'Right'
                    ORDER BY created_at DESC
                    LIMIT 1
                """, (athlete_uuid, session_date))
                right_data = cur.fetchone()
                
                return (left_data, right_data)
            else:
                cur.execute(f"""
                    SELECT {', '.join(columns)}
                    FROM public.{pg_table}
                    WHERE athlete_uuid = %s AND session_date = %s
                    ORDER BY created_at DESC
                    LIMIT 1
                """, (athlete_uuid, session_date))
                return cur.fetchone()
    except Exception as e:
        print(f"Error querying client data for {movement_type}: {e}")
        return None


def query_reference_data(conn, movement_type: str, columns: List[str]) -> pd.DataFrame:
    """Query all reference data from PostgreSQL (for percentile calculations)."""
    pg_table = MOVEMENT_TO_PG_TABLE.get(movement_type)
    if not pg_table:
        return pd.DataFrame()
    
    try:
        query = f"""
            SELECT {', '.join(columns)}
            FROM public.{pg_table}
            WHERE {columns[0]} IS NOT NULL
        """
        return pd.read_sql_query(query, conn)
    except Exception as e:
        print(f"Error querying reference data for {movement_type}: {e}")
        return pd.DataFrame()


def generate_report(athlete_uuid: str, athlete_name: str, session_date: str,
                   output_dir: str, logo_path: Optional[str] = None,
                   power_files_dir: Optional[str] = None) -> str:
    """
    Generate Athletic Screen report comparing client to "All" reference data.
    
    Args:
        athlete_uuid: UUID of the athlete
        athlete_name: Name of the athlete
        session_date: Session date (YYYY-MM-DD)
        output_dir: Directory for output report
        logo_path: Optional path to logo image
        power_files_dir: Directory containing power files
    
    Returns:
        Path to generated report file
    """
    conn = get_warehouse_connection()
    
    try:
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        # Build output filename
        parts = athlete_name.split(', ')
        client_name_rev = f"{parts[1]}_{parts[0]}" if len(parts) == 2 else athlete_name
        base_name = f"Athletic_Report_{client_name_rev}_{session_date}_All"
        output_filename = os.path.join(output_dir, base_name + ".docx")
        
        # Auto-increment if file exists
        counter = 1
        while os.path.exists(output_filename):
            output_filename = os.path.join(output_dir, f"{base_name}_{counter}.docx")
            counter += 1
        
        # Create document
        doc = Document()
        
        # Add logo if provided
        if logo_path and os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(4.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add header
        doc.add_paragraph(f"Player's Name: {athlete_name}")
        doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}")
        
        # Create temporary directory for images
        with tempfile.TemporaryDirectory() as tmpdirname:
            movements = ['CMJ', 'DJ', 'PPU', 'SLV', 'NMT']
            
            for movement in movements:
                doc.add_paragraph(f"{movement} Report", style='Title')
                doc.add_paragraph(f"This section includes percentile reports and comparisons for {movement}.", style='Heading 2')
                
                if movement == 'CMJ':
                    _process_cmj_section(doc, conn, athlete_uuid, session_date, tmpdirname, power_files_dir)
                elif movement == 'DJ':
                    _process_dj_section(doc, conn, athlete_uuid, session_date, tmpdirname, power_files_dir)
                elif movement == 'PPU':
                    _process_ppu_section(doc, conn, athlete_uuid, session_date, tmpdirname, power_files_dir)
                elif movement == 'SLV':
                    _process_slv_section(doc, conn, athlete_uuid, session_date, tmpdirname, power_files_dir)
                elif movement == 'NMT':
                    _process_nmt_section(doc, conn, athlete_uuid, session_date, tmpdirname)
        
        # Save document
        doc.save(output_filename)
        print(f"Report saved: {output_filename}")
        return output_filename
        
    finally:
        conn.close()


def _process_cmj_section(doc: Document, conn, athlete_uuid: str, session_date: str,
                         tmpdirname: str, power_files_dir: Optional[str]):
    """Process CMJ section of report."""
    variables = ['jh_in', 'pp_forceplate', 'pp_w_per_kg', 'force_at_pp', 'vel_at_pp']
    columns = variables  # PostgreSQL uses lowercase column names
    
    # Query client data
    client_data = query_client_data(conn, athlete_uuid, session_date, 'CMJ', columns)
    if not client_data:
        doc.add_paragraph("No CMJ data available for this session.")
        return
    
    # Query reference data
    ref_df = query_reference_data(conn, 'CMJ', columns)
    if ref_df.empty:
        doc.add_paragraph("No reference data available for comparison.")
        return
    
    # Power analysis (if files available)
    if power_files_dir:
        cmj_files = find_power_files("CMJ", power_files_dir)
        if cmj_files:
            try:
                cmj_traces = [load_power_txt(pf) for pf in cmj_files]
                # TODO: Add power analysis section (requires additional functions from notebook)
            except Exception as e:
                print(f"Warning: Could not load power files: {e}")
    
    # Generate bar graphs
    for i, var in enumerate(variables):
        formatted_var = var.replace('_', ' ').title()
        doc.add_paragraph(f"{formatted_var} Comparison", style='Heading 2')
        
        client_val = float(client_data[i]) if client_data[i] is not None else 0.0
        ref_col = ref_df[columns[i]].dropna().values
        
        if len(ref_col) > 0:
            bar_image = generate_bar_graph(columns[i].upper(), client_val, ref_col,
                                          f'{formatted_var} Comparison', tmpdirname,
                                          conn, athlete_uuid, session_date, 'CMJ')
            doc.add_picture(bar_image, width=Inches(6))
    
    # Scatter plot
    if client_data[3] is not None and client_data[4] is not None:
        doc.add_paragraph("Force vs. Velocity Scatter Plot", style='Heading 2')
        client_dict = {
            'Force_at_PP': float(client_data[3]),
            'Vel_at_PP': float(client_data[4])
        }
        ref_scatter = ref_df[['force_at_pp', 'vel_at_pp']].rename(columns={
            'force_at_pp': 'Force_at_PP',
            'vel_at_pp': 'Vel_at_PP'
        })
        scatter_image = generate_scatter_plot(client_dict, ref_scatter,
                                            'Force_at_PP', 'Vel_at_PP',
                                            'CMJ: Force vs. Velocity', tmpdirname)
        doc.add_picture(scatter_image, width=Inches(6))


def _process_dj_section(doc: Document, conn, athlete_uuid: str, session_date: str,
                       tmpdirname: str, power_files_dir: Optional[str]):
    """Process DJ section of report."""
    variables = ['jh_in', 'pp_forceplate', 'pp_w_per_kg', 'force_at_pp', 'vel_at_pp', 'ct', 'rsi']
    columns = variables  # PostgreSQL uses lowercase column names
    
    client_data = query_client_data(conn, athlete_uuid, session_date, 'DJ', columns)
    if not client_data:
        doc.add_paragraph("No DJ data available for this session.")
        return
    
    ref_df = query_reference_data(conn, 'DJ', columns)
    if ref_df.empty:
        doc.add_paragraph("No reference data available for comparison.")
        return
    
    # Generate bar graphs
    for i, var in enumerate(variables):
        formatted_var = var.replace('_', ' ').title()
        doc.add_paragraph(f"{formatted_var} Comparison", style='Heading 2')
        
        client_val = float(client_data[i]) if client_data[i] is not None else 0.0
        ref_col = ref_df[columns[i]].dropna().values
        
        if len(ref_col) > 0:
            bar_image = generate_bar_graph(columns[i].upper(), client_val, ref_col,
                                          f'{formatted_var} Comparison', tmpdirname,
                                          conn, athlete_uuid, session_date, 'DJ')
            doc.add_picture(bar_image, width=Inches(6))
    
    # Scatter plot
    if client_data[3] is not None and client_data[4] is not None:
        doc.add_paragraph("Force vs. Velocity Scatter Plot", style='Heading 2')
        client_dict = {
            'Force_at_PP': float(client_data[3]),
            'Vel_at_PP': float(client_data[4])
        }
        ref_scatter = ref_df[['force_at_pp', 'vel_at_pp']].rename(columns={
            'force_at_pp': 'Force_at_PP',
            'vel_at_pp': 'Vel_at_PP'
        })
        scatter_image = generate_scatter_plot(client_dict, ref_scatter,
                                            'Force_at_PP', 'Vel_at_PP',
                                            'DJ: Force vs. Velocity', tmpdirname)
        doc.add_picture(scatter_image, width=Inches(6))


def _process_ppu_section(doc: Document, conn, athlete_uuid: str, session_date: str,
                        tmpdirname: str, power_files_dir: Optional[str]):
    """Process PPU section of report."""
    variables = ['jh_in', 'pp_forceplate', 'pp_w_per_kg', 'force_at_pp', 'vel_at_pp']
    columns = variables  # PostgreSQL uses lowercase column names
    
    client_data = query_client_data(conn, athlete_uuid, session_date, 'PPU', columns)
    if not client_data:
        doc.add_paragraph("No PPU data available for this session.")
        return
    
    ref_df = query_reference_data(conn, 'PPU', columns)
    if ref_df.empty:
        doc.add_paragraph("No reference data available for comparison.")
        return
    
    # Generate bar graphs
    for i, var in enumerate(variables):
        formatted_var = var.replace('_', ' ').title()
        doc.add_paragraph(f"{formatted_var} Comparison", style='Heading 2')
        
        client_val = float(client_data[i]) if client_data[i] is not None else 0.0
        ref_col = ref_df[columns[i]].dropna().values
        
        if len(ref_col) > 0:
            bar_image = generate_bar_graph(columns[i].upper(), client_val, ref_col,
                                          f'{formatted_var} Comparison', tmpdirname,
                                          conn, athlete_uuid, session_date, 'PPU')
            doc.add_picture(bar_image, width=Inches(6))
    
    # Scatter plot
    if client_data[3] is not None and client_data[4] is not None:
        doc.add_paragraph("Force vs. Velocity Scatter Plot", style='Heading 2')
        client_dict = {
            'Force_at_PP': float(client_data[3]),
            'Vel_at_PP': float(client_data[4])
        }
        ref_scatter = ref_df[['force_at_pp', 'vel_at_pp']].rename(columns={
            'force_at_pp': 'Force_at_PP',
            'vel_at_pp': 'Vel_at_PP'
        })
        scatter_image = generate_scatter_plot(client_dict, ref_scatter,
                                            'Force_at_PP', 'Vel_at_PP',
                                            'PPU: Force vs. Velocity', tmpdirname)
        doc.add_picture(scatter_image, width=Inches(6))


def _process_slv_section(doc: Document, conn, athlete_uuid: str, session_date: str,
                        tmpdirname: str, power_files_dir: Optional[str]):
    """Process SLV section of report (handles Left/Right sides)."""
    variables = ['jh_in', 'pp_forceplate', 'pp_w_per_kg', 'force_at_pp', 'vel_at_pp']
    columns = variables  # PostgreSQL uses lowercase column names
    
    client_data = query_client_data(conn, athlete_uuid, session_date, 'SLV', columns)
    if not client_data or len(client_data) != 2:
        doc.add_paragraph("No SLV data available for this session (both sides required).")
        return
    
    left_data, right_data = client_data
    if not left_data or not right_data:
        doc.add_paragraph("Incomplete SLV data (both Left and Right sides required).")
        return
    
    ref_df = query_reference_data(conn, 'SLV', columns)
    if ref_df.empty:
        doc.add_paragraph("No reference data available for comparison.")
        return
    
    # Generate histograms for each variable (Left vs Right)
    for i, var in enumerate(variables):
        formatted_var = var.replace('_', ' ').title()
        doc.add_paragraph(f"{formatted_var} Comparison (Left vs Right)", style='Heading 2')
        
        left_val = float(left_data[i]) if left_data[i] is not None else 0.0
        right_val = float(right_data[i]) if right_data[i] is not None else 0.0
        ref_col = ref_df[columns[i]].dropna().values
        
        if len(ref_col) > 0:
            # Use generate_bar_graph but we'll need to modify for SLV
            # For now, create separate graphs for left and right
            left_image = generate_bar_graph(columns[i].upper(), left_val, ref_col,
                                          f'{formatted_var} Left', tmpdirname,
                                          conn, athlete_uuid, session_date, 'SLV')
            doc.add_picture(left_image, width=Inches(6))
            
            right_image = generate_bar_graph(columns[i].upper(), right_val, ref_col,
                                           f'{formatted_var} Right', tmpdirname,
                                           conn, athlete_uuid, session_date, 'SLV')
            doc.add_picture(right_image, width=Inches(6))
    
    # Scatter plot (both sides)
    if (left_data[3] is not None and left_data[4] is not None and
        right_data[3] is not None and right_data[4] is not None):
        doc.add_paragraph("Force vs. Velocity Scatter Plot", style='Heading 2')
        client_dict = {
            'Force_at_PP': [float(left_data[3]), float(right_data[3])],
            'Vel_at_PP': [float(left_data[4]), float(right_data[4])]
        }
        ref_scatter = ref_df[['force_at_pp', 'vel_at_pp']].rename(columns={
            'force_at_pp': 'Force_at_PP',
            'vel_at_pp': 'Vel_at_PP'
        })
        scatter_image = generate_scatter_plot(client_dict, ref_scatter,
                                            'Force_at_PP', 'Vel_at_PP',
                                            'SLV: Force vs. Velocity', tmpdirname)
        doc.add_picture(scatter_image, width=Inches(6))


def _process_nmt_section(doc: Document, conn, athlete_uuid: str, session_date: str,
                        tmpdirname: str):
    """Process NMT section of report."""
    variables = ['num_taps_10s', 'num_taps_20s', 'num_taps_30s', 'num_taps']
    columns = variables  # PostgreSQL uses lowercase column names
    
    client_data = query_client_data(conn, athlete_uuid, session_date, 'NMT', columns)
    if not client_data:
        doc.add_paragraph("No NMT data available for this session.")
        return
    
    ref_df = query_reference_data(conn, 'NMT', columns)
    if ref_df.empty:
        doc.add_paragraph("No reference data available for comparison.")
        return
    
    # Generate bar graphs (typically only 10s is used)
    for i, var in enumerate(variables):
        formatted_var = var.replace('_', ' ').title()
        doc.add_paragraph(f"{formatted_var} Comparison", style='Heading 2')
        
        client_val = float(client_data[i]) if client_data[i] is not None else 0.0
        ref_col = ref_df[columns[i]].dropna().values
        
        if len(ref_col) > 0:
            bar_image = generate_bar_graph(columns[i].upper(), client_val, ref_col,
                                          f'{formatted_var} Comparison', tmpdirname,
                                          conn, athlete_uuid, session_date, 'NMT')
            doc.add_picture(bar_image, width=Inches(6))

